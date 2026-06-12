from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
import re
import time

import requests
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from pypdf import PdfReader
from tqdm import tqdm


MODEL = "ali6parmak/hy-mt1.5:1.8b"
OLLAMA_URL = "http://localhost:11434/api/chat"
INPUT_DIR = Path("input")
OUTPUT_DIR = Path("output")
TARGET_LANGUAGE = "中文"
OUTPUT_SUFFIX = "中英对照"
MAX_CHARS = 1200
MAX_WORKERS = 4          # 并发线程数，可根据显存/CPU往上调，建议先从4开始
MERGE_MIN_CHARS = 80     # 短于此字符数的句子会被合并，减少 API 调用次数
SENTENCE_END_RE = re.compile(r"(?<=[。！？.!?；;])\s+|(?<=[。！？.!?；;])")


def translate_text(text, target_language=TARGET_LANGUAGE, retries=3):
    text = text.strip()
    if not text:
        return ""

    payload = {
        "model": MODEL,
        "messages": [
            {
                "role": "user",
                "content": f"将以下文本翻译为{target_language}。只输出翻译结果，不要解释，不要加标题：\n{text}",
            }
        ],
        "stream": False,
    }

    last_err = None
    for _ in range(retries):
        try:
            r = requests.post(OLLAMA_URL, json=payload, timeout=600)
            r.raise_for_status()
            data = r.json()
            return data["message"]["content"].strip()
        except Exception as e:
            last_err = e
            time.sleep(2)

    raise last_err


def split_chunks(text, max_chars=MAX_CHARS):
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    paragraphs = re.split(r"(\n\s*\n)", text)
    chunks = []
    buf = ""

    for part in paragraphs:
        if len(buf) + len(part) > max_chars and buf.strip():
            chunks.append(buf.strip())
            buf = part
        else:
            buf += part

    if buf.strip():
        chunks.append(buf.strip())

    return chunks


def split_sentences(text):
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    sentences = []

    for paragraph in re.split(r"\n\s*\n", text):
        paragraph = " ".join(line.strip() for line in paragraph.splitlines() if line.strip())
        if not paragraph:
            continue

        for sentence in SENTENCE_END_RE.split(paragraph):
            sentence = sentence.strip()
            if not sentence:
                continue
            sentences.extend(split_chunks(sentence))

    return sentences


def merge_short_sentences(sentences, min_chars=MERGE_MIN_CHARS):
    """将过短的句子合并，减少 API 调用次数"""
    merged, buf = [], ""
    for s in sentences:
        buf = (buf + " " + s).strip() if buf else s
        if len(buf) >= min_chars:
            merged.append(buf)
            buf = ""
    if buf:
        merged.append(buf)
    return merged


def translate_bilingual_text(text, target_language=TARGET_LANGUAGE):
    sentences = split_sentences(text)
    if not sentences:
        return ""

    sentences = merge_short_sentences(sentences)  # 合并短句，减少请求数
    separator = "\n\n" + "=" * 40 + "\n\n"
    results = [None] * len(sentences)

    def _job(idx, sentence):
        return idx, translate_text(sentence, target_language)

    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as pool:
        futures = {pool.submit(_job, i, s): i for i, s in enumerate(sentences)}
        for future in tqdm(
            as_completed(futures),
            total=len(sentences),
            desc=f"Translating ({len(sentences)} sentences)",
            leave=False,
        ):
            idx, translated = future.result()
            results[idx] = f"{sentences[idx]}\n{translated}"

    return separator.join(results)


def read_txt(path):
    for enc in ("utf-8", "utf-8-sig", "gbk", "gb18030"):
        try:
            return path.read_text(encoding=enc)
        except Exception:
            pass

    return path.read_text(encoding="utf-8", errors="ignore")


def write_txt(path, text):
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")


def read_docx(path):
    doc = Document(str(path))
    lines = []

    for p in doc.paragraphs:
        if p.text.strip():
            lines.append(p.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            if any(row_text):
                lines.append("\t".join(row_text))

    return "\n".join(lines)


def set_run_font(run):
    run.font.name = "宋体"
    run.font.size = Pt(12)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), "宋体")


def write_docx(path, text):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)

    r_fonts = style.element.xpath("w:rPr/w:rFonts")
    if r_fonts:
        r_fonts[0].set(qn("w:eastAsia"), "宋体")
    else:
        r_pr = style.element.get_or_add_rPr()
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:eastAsia"), "宋体")
        r_pr.append(r_fonts)

    for para_text in text.split("\n\n"):
        paragraph = doc.add_paragraph(para_text)
        for run in paragraph.runs:
            set_run_font(run)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def read_pdf(path):
    reader = PdfReader(str(path))
    pages = []

    for page in reader.pages:
        pages.append(page.extract_text() or "")

    return "\n\n".join(pages)


def process_file(file_path, target_language=TARGET_LANGUAGE):
    ext = file_path.suffix.lower()

    if ext in [".txt", ".md", ".log"]:
        text = read_txt(file_path)
        bilingual = translate_bilingual_text(text, target_language)
        out_path = OUTPUT_DIR / f"{file_path.stem}.{OUTPUT_SUFFIX}.txt"
        write_txt(out_path, bilingual)
        return out_path

    if ext == ".docx":
        text = read_docx(file_path)
        bilingual = translate_bilingual_text(text, target_language)
        out_path = OUTPUT_DIR / f"{file_path.stem}.{OUTPUT_SUFFIX}.docx"
        write_docx(out_path, bilingual)
        return out_path

    if ext == ".pdf":
        text = read_pdf(file_path)
        bilingual = translate_bilingual_text(text, target_language)
        out_path = OUTPUT_DIR / f"{file_path.stem}.{OUTPUT_SUFFIX}.docx"
        write_docx(out_path, bilingual)
        return out_path

    return None


def batch_translate():
    INPUT_DIR.mkdir(exist_ok=True)
    OUTPUT_DIR.mkdir(exist_ok=True)

    files = [p for p in INPUT_DIR.rglob("*") if p.is_file()]
    supported = [
        p for p in files if p.suffix.lower() in [".txt", ".md", ".log", ".docx", ".pdf"]
    ]

    if not supported:
        print("input 文件夹里没有可翻译文件。")
        return

    for file_path in tqdm(supported, desc="Translating"):
        try:
            out = process_file(file_path)
            if out:
                print(f"OK: {file_path.name} -> {out}")
            else:
                print(f"SKIP: {file_path.name}")
        except Exception as e:
            print(f"FAIL: {file_path.name} -> {e}")


if __name__ == "__main__":
    batch_translate()
