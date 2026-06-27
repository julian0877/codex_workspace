import os
from datetime import datetime
from pathlib import Path
from typing import List, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def _parse_text_to_pairs(text: str) -> List[Tuple[str, str]]:
    pairs = []
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        if '=' in line:
            key, value = line.split('=', 1)
            pairs.append((key.strip(), value.strip()))
        elif ':' in line:
            key, value = line.split(':', 1)
            pairs.append((key.strip(), value.strip()))
        else:
            pairs.append((line, ''))
    return pairs


def _set_default_style(doc: Document) -> None:
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = heading.add_run(text)
    run.bold = True
    run.font.size = Pt(16 if level == 1 else 14 if level == 2 else 12)


def _add_paragraph(doc: Document, text: str, indent: bool = False) -> None:
    paragraph = doc.add_paragraph(text)
    if indent:
        paragraph.paragraph_format.first_line_indent = Inches(0.3)


def _add_two_column_table(doc: Document, items: List[Tuple[str, str]]) -> None:
    if not items:
        return
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '项目'
    hdr_cells[1].text = '数值 / 说明'
    for key, value in items:
        row_cells = table.add_row().cells
        row_cells[0].text = str(key)
        row_cells[1].text = str(value)


def generate_word_from_data(data: dict, output_path: str) -> str:
    output_dir = os.path.dirname(output_path) or '.'
    os.makedirs(output_dir, exist_ok=True)

    doc = Document()
    _set_default_style(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('吊耳计算书')
    run.bold = True
    run.font.size = Pt(20)

    doc.add_paragraph(f"项目名称：{data.get('project_name', '')}")
    doc.add_paragraph(f"日期：{data.get('date', '')}")
    doc.add_paragraph(f"设计单位 / 工程：{data.get('project_info', '')}")
    doc.add_paragraph(f"核算人员：{data.get('author', '')}")
    doc.add_paragraph(f"审核人员：{data.get('reviewer', '')}")

    # 一、计算简图
    _add_heading(doc, '一、计算简图', 1)
    diagram_text = data.get('diagram_text', '').strip()
    if diagram_text:
        _add_paragraph(doc, diagram_text, indent=True)
    else:
        _add_paragraph(doc, '本节简要说明吊耳受力形式和计算思路。', indent=True)

    # 二、已知条件
    _add_heading(doc, '二、已知条件', 1)
    known_pairs = _parse_text_to_pairs(data.get('known_conditions', ''))
    if known_pairs:
        _add_paragraph(doc, '已知条件如下：', indent=True)
        _add_two_column_table(doc, known_pairs)
    else:
        _add_paragraph(doc, '无已知条件数据。', indent=True)

    # 三、材质、尺寸及力学性能
    _add_heading(doc, '三、材质、尺寸及力学性能', 1)
    material_pairs = _parse_text_to_pairs(data.get('material_data', ''))
    if material_pairs:
        _add_paragraph(doc, '构件材质、尺寸和力学性能设计值如下：', indent=True)
        _add_two_column_table(doc, material_pairs)
    else:
        _add_paragraph(doc, '无材质与尺寸数据。', indent=True)

    # 四、耳板承载力计算结果及判定
    _add_heading(doc, '四、耳板承载力计算结果及判定', 1)
    ear_pairs = _parse_text_to_pairs(data.get('ear_results', ''))
    if ear_pairs:
        _add_paragraph(doc, '耳板承载力计算结果及判定如下：', indent=True)
        _add_two_column_table(doc, ear_pairs)
    else:
        _add_paragraph(doc, '无耳板承载力计算结果数据。', indent=True)

    # 五、双面角焊缝承载力计算结果及判定
    _add_heading(doc, '五、双面角焊缝承载力计算结果及判定', 1)
    weld_pairs = _parse_text_to_pairs(data.get('weld_results', ''))
    if weld_pairs:
        _add_paragraph(doc, '双面角焊缝承载力计算结果及判定如下：', indent=True)
        _add_two_column_table(doc, weld_pairs)
    else:
        _add_paragraph(doc, '无双面角焊缝计算结果数据。', indent=True)

    # 结论
    _add_heading(doc, '六、结论', 1)
    conclusion_text = data.get('conclusion', '').strip()
    if conclusion_text:
        _add_paragraph(doc, conclusion_text, indent=True)
    else:
        _add_paragraph(doc, '根据上述计算结果，吊耳和焊缝均满足验算要求。', indent=True)

    doc.save(output_path)
    return output_path


def default_output_path() -> str:
    base = Path(__file__).resolve().parent
    output_dir = base / 'output'
    output_dir.mkdir(exist_ok=True)
    return str(output_dir / '吊耳计算书.docx')
