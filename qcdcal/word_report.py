from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


@dataclass(frozen=True)
class CalculationReportSection:
    title: str
    inputs: dict[str, tuple[str, str]]
    formulas: list[str]
    results: dict[str, tuple[str, str]]


def export_word_report(
    output_path: str | Path,
    sections: list[CalculationReportSection],
) -> Path:
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _apply_document_style(doc)
    _add_title(doc)

    for index, section in enumerate(sections, start=1):
        doc.add_heading(f"{index}. {section.title}", level=1)
        _add_key_value_table(doc, "输入参数", section.inputs)
        _add_formula_list(doc, section.formulas)
        _add_key_value_table(doc, "计算结果", section.results)

    doc.save(path)
    return path


def _apply_document_style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(0x2E, 0x74, 0xB5)),
        ("Heading 2", 13, RGBColor(0x2E, 0x74, 0xB5)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def _add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("吊装相关验算计算书")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x4D, 0x78)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")


def _add_formula_list(doc: Document, formulas: list[str]) -> None:
    doc.add_heading("计算公式", level=2)
    if not formulas:
        doc.add_paragraph("无")
        return

    for formula in formulas:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(formula)


def _add_key_value_table(
    doc: Document,
    title: str,
    values: dict[str, tuple[str, str]],
) -> None:
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    _set_table_width(table, [Inches(3.8), Inches(1.35), Inches(1.1)])

    headers = ["项目", "数值", "单位"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        _shade_cell(cell, "F2F4F7")
        _set_cell_text_bold(cell)

    if values:
        for label, (value, unit) in values.items():
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = value
            cells[2].text = unit
            cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for cell in cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    else:
        cells = table.add_row().cells
        cells[0].text = "无"
        cells[1].text = "-"
        cells[2].text = "-"

    doc.add_paragraph()


def _set_table_width(table, widths) -> None:
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), fill)
    tc_pr.append(shade)


def _set_cell_text_bold(cell) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
