import tempfile
import unittest
from pathlib import Path

from docx import Document

from lifting_calculations import (
    overturning_calculation,
    outrigger_reaction_calculation,
)
from word_report import CalculationReportSection, export_word_report


class WordReportTests(unittest.TestCase):
    def test_exports_word_report_with_two_calculation_sections(self):
        overturning_inputs = {
            "汽车吊自重 G": ("72", "t"),
            "起吊构件最大重量 Q": ("27", "t"),
        }
        overturning_results = {
            "风荷载 W": ("52.920", "kN"),
            "倾覆边合力矩 M": ("6941.487", "kN.m"),
            "结果判断": ("满足要求", ""),
        }
        outrigger_inputs = {
            "汽车吊吊臂自重 G3": ("10", "t"),
            "路基箱面积 S": ("4", "m2"),
        }
        outrigger_results = {
            "最大支腿反力": ("245.366", "kN"),
            "最大地面压强 P": ("61.341", "kPa"),
        }

        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = Path(temp_dir) / "吊装计算书.docx"
            export_word_report(
                output_path,
                [
                    CalculationReportSection(
                        title="汽车吊抗倾覆验算",
                        inputs=overturning_inputs,
                        formulas=[
                            "W=20%*Q*g",
                            "M=KG*MG+KQ*MQ+KW*MW",
                        ],
                        results=overturning_results,
                    ),
                    CalculationReportSection(
                        title="汽车吊支腿反力计算",
                        inputs=outrigger_inputs,
                        formulas=[
                            "F=(G0+G1+G2+G3)*g",
                            "P=N/S",
                        ],
                        results=outrigger_results,
                    ),
                ],
            )

            self.assertTrue(output_path.exists())
            text = "\n".join(
                paragraph.text for paragraph in Document(output_path).paragraphs
            )
            table_text = "\n".join(
                cell.text
                for table in Document(output_path).tables
                for row in table.rows
                for cell in row.cells
            )
            all_text = text + "\n" + table_text

        self.assertIn("吊装相关验算计算书", all_text)
        self.assertIn("汽车吊抗倾覆验算", all_text)
        self.assertIn("汽车吊支腿反力计算", all_text)
        self.assertIn("倾覆边合力矩 M", all_text)
        self.assertIn("最大地面压强 P", all_text)
        self.assertIn("满足要求", all_text)

    def test_report_accepts_real_calculation_outputs(self):
        overturning = overturning_calculation(
            crane_weight=72,
            lifted_weight=27,
            counterweight=100,
            working_radius=20,
            wind_height=22,
            crane_center_to_tip=4.3,
            counterweight_to_center=5.75,
            gravity=9.8,
            crane_weight_factor=1,
            lifted_weight_factor=1.15,
            wind_factor=1,
        )
        outrigger = outrigger_reaction_calculation(
            boom_weight=10,
            boom_center_distance=4,
            crane_weight_without_boom=31,
            lifted_weight=6.02,
            counterweight=8,
            working_radius=10,
            longitudinal_distance=5.92,
            transverse_distance=6.9,
            counterweight_to_center=0,
            gravity=9.8,
            center_to_rear_outrigger=2.9,
            ground_box_area=4,
        )

        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = Path(temp_dir) / "sample.docx"
            export_word_report(
                output_path,
                [
                    CalculationReportSection(
                        title="汽车吊抗倾覆验算",
                        inputs={},
                        formulas=[],
                        results={"倾覆边合力矩 M": (f"{overturning['resultant_moment']:.3f}", "kN.m")},
                    ),
                    CalculationReportSection(
                        title="汽车吊支腿反力计算",
                        inputs={},
                        formulas=[],
                        results={"最大地面压强 P": (f"{outrigger['max_ground_pressure']:.3f}", "kPa")},
                    ),
                ],
            )

            self.assertGreater(output_path.stat().st_size, 0)


if __name__ == "__main__":
    unittest.main()
