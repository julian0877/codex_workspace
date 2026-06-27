import os
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document
from openpyxl import Workbook

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from cal.generate_word_calculation_book import generate_word_from_excel


class GenerateWordCalculationBookTest(unittest.TestCase):
    def test_generate_docx_from_excel(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            excel_path = os.path.join(tmpdir, 'sample.xlsx')
            output_path = os.path.join(tmpdir, 'sample.docx')

            wb = Workbook()
            ws = wb.active
            ws.title = '吊耳计算'
            ws['A1'] = '项目'
            ws['B1'] = '数值'
            ws['A2'] = '吊耳型号'
            ws['B2'] = '双吊耳'
            ws['A3'] = '吊装重量'
            ws['B3'] = '5t'
            ws['A4'] = '结论'
            ws['B4'] = '满足要求'
            wb.save(excel_path)

            result = generate_word_from_excel(excel_path, output_path)

            self.assertTrue(result)
            self.assertTrue(os.path.exists(output_path))

            doc = Document(output_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = '\n'.join(paragraphs)
            self.assertIn('吊耳计算书', text)
            self.assertIn('一、计算简图', text)
            self.assertIn('二、已知条件', text)
            self.assertIn('三、材质、尺寸及力学性能', text)
            self.assertIn('四、耳板承载力计算结果及判定', text)
            self.assertIn('五、双面角焊缝承载力计算结果及判定', text)
            self.assertTrue(len(doc.tables) >= 1)


if __name__ == '__main__':
    unittest.main()
