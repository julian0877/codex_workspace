from pathlib import Path

import pytest
from docx import Document

from tender_formatter.word_adapter import (
    WordAdapter,
    WordAutomationError,
    WordSettings,
)


class FakeCollection:
    def __init__(self, items):
        self._items = items

    @property
    def Count(self):
        return len(self._items)

    def __call__(self, index):
        return self._items[index - 1]


class FakeUpdatable:
    def __init__(self, document, name, fail_on=None):
        self.document = document
        self.name = name
        self.fail_on = fail_on

    def Update(self):
        if self.fail_on == self.name:
            raise RuntimeError("COM failure")
        setattr(self.document, f"{self.name.lower()}_updated", True)


class FakePageNumbers:
    RestartNumberingAtSection = False
    StartingNumber = 0


class FakeHeaderFooter:
    LinkToPrevious = True

    def __init__(self):
        self.PageNumbers = FakePageNumbers()


class FakeSection:
    def __init__(self):
        self.PageSetup = type(
            "PageSetup", (), {"DifferentFirstPageHeaderFooter": False, "OddAndEvenPagesHeaderFooter": False}
        )()
        self.Headers = FakeCollection([FakeHeaderFooter(), FakeHeaderFooter(), FakeHeaderFooter()])
        self.Footers = FakeCollection([FakeHeaderFooter(), FakeHeaderFooter(), FakeHeaderFooter()])


class FakeDocument:
    def __init__(self, fail_on=None):
        self.fields_updated = False
        self.toc_updated = False
        self.saved = False
        self.closed = False
        self.Fields = FakeUpdatable(self, "Fields", fail_on)
        self.TablesOfContents = FakeCollection([FakeUpdatable(self, "Toc", fail_on)])
        self.Sections = FakeCollection([FakeSection()])

    def Repaginate(self):
        return None

    def Save(self):
        self.saved = True

    def Close(self, _save_changes=False):
        self.closed = True


class FakeDocuments:
    def __init__(self, document):
        self.document = document
        self.opened_path = None

    def Open(self, path):
        self.opened_path = path
        return self.document


class FakeWordApplication:
    def __init__(self, fail_on=None):
        self.Visible = True
        self.DisplayAlerts = 1
        self.opened_document = FakeDocument(fail_on)
        self.Documents = FakeDocuments(self.opened_document)
        self.quit_called = False

    def Quit(self):
        self.quit_called = True


def test_finalize_uses_private_word_instance_and_always_quits(tmp_path):
    app = FakeWordApplication()
    adapter = WordAdapter(dispatch_ex=lambda _name: app)

    adapter.finalize(
        tmp_path / "out.docx", WordSettings(update_toc=True, body_page_start=1)
    )

    assert app.Visible is False
    assert app.opened_document.fields_updated
    assert app.opened_document.toc_updated
    assert app.opened_document.saved
    assert app.opened_document.closed
    assert app.quit_called


def test_finalize_quits_private_instance_on_failure(tmp_path):
    app = FakeWordApplication(fail_on="Fields")

    with pytest.raises(WordAutomationError, match="Word 自动化失败"):
        WordAdapter(dispatch_ex=lambda _name: app).finalize(
            tmp_path / "out.docx", WordSettings()
        )

    assert app.opened_document.closed
    assert app.quit_called


@pytest.mark.word_integration
def test_real_word_assembles_template_cover_toc_content_and_page_number(tmp_path):
    template = tmp_path / "template.docx"
    template_document = Document()
    template_document.add_paragraph("{{项目名称}}")
    template_document.add_paragraph("{{目录}}")
    template_document.save(template)
    content = tmp_path / "content.docx"
    content_document = Document()
    content_document.add_paragraph("第一章 施工部署", style="Heading 1")
    content_document.add_paragraph("正文内容")
    content_document.save(content)
    output = tmp_path / "assembled.docx"

    WordAdapter().assemble(
        content,
        output,
        template,
        WordSettings(),
        {"项目名称": "厂房项目"},
    )

    assembled = Document(output)
    text = "\n".join(paragraph.text for paragraph in assembled.paragraphs)
    xml = assembled._element.xml
    footer_xml = "".join(
        section.footer._element.xml for section in assembled.sections
    )
    assert "厂房项目" in text
    assert "第一章 施工部署" in text
    assert "{{目录}}" not in text
    assert 'w:instrText xml:space="preserve"> TOC ' in xml
    assert "> PAGE <" in footer_xml
