from docx import Document

from tender_formatter.cover import find_cover_fields, replace_cover_fields
from tests.helpers import make_cover_docx


def test_cover_fields_are_detected_and_replaced(tmp_path):
    path = make_cover_docx(
        tmp_path / "cover.docx", "{{项目名称}}\n{{投标单位}}\n{{日期}}"
    )
    assert find_cover_fields(path) == {"项目名称", "投标单位", "日期"}
    document = Document(path)

    missing = replace_cover_fields(
        document,
        {
            "项目名称": "厂房项目",
            "投标单位": "建设公司",
            "日期": "2026年7月",
        },
    )

    assert missing == []
    assert "{{" not in "\n".join(p.text for p in document.paragraphs)


def test_missing_cover_value_is_reported(tmp_path):
    document = Document(
        make_cover_docx(tmp_path / "cover.docx", "{{项目名称}}\n{{日期}}")
    )

    missing = replace_cover_fields(document, {"项目名称": "厂房项目"})

    assert missing == ["日期"]
    assert "{{日期}}" in "\n".join(p.text for p in document.paragraphs)


def test_placeholder_split_across_runs_is_replaced(tmp_path):
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("{{项目")
    paragraph.add_run("名称}}")

    assert replace_cover_fields(document, {"项目名称": "厂房项目"}) == []
    assert paragraph.text == "厂房项目"
