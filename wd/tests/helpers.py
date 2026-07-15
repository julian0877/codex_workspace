from pathlib import Path
import base64

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def _set_style_font(style, east_asia_font: str, size_pt: float, bold: bool) -> None:
    style.font.name = east_asia_font
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)


def make_docx_with_styles(path: Path) -> Path:
    document = Document()
    _set_style_font(document.styles["Normal"], "宋体", 12, False)
    _set_style_font(document.styles["Heading 1"], "黑体", 16, True)
    _set_style_font(document.styles["Heading 2"], "黑体", 14, True)
    _set_style_font(document.styles["Heading 3"], "黑体", 12, True)
    document.save(path)
    return path


def make_mixed_docx(path: Path) -> Path:
    document = Document()
    document.add_paragraph("第一章 总体方案", style="Heading 1")
    document.add_paragraph("1.1 编制依据", style="Heading 2")
    document.add_paragraph("正文内容")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "项目"
    table.cell(0, 1).text = "内容"
    image_path = path.with_suffix(".png")
    image_path.write_bytes(
        base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
        )
    )
    document.add_picture(str(image_path), width=Cm(20))
    document.save(path)
    image_path.unlink()
    return path
