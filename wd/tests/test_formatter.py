from docx import Document
from docx.oxml.ns import qn

from tender_formatter.domain import FormatOperation, FormatPlan
from tender_formatter.formatter import execute_docx_plan
from tests.helpers import make_mixed_docx


def make_plan(tmp_path):
    source = make_mixed_docx(tmp_path / "source.docx")
    return FormatPlan(
        source=source,
        output=tmp_path / "source_已格式化.docx",
        operations=[
            FormatOperation(
                kind="apply_paragraph",
                target=0,
                parameters={"kind": "heading", "level": 1},
            ),
            FormatOperation(
                kind="apply_paragraph",
                target=1,
                parameters={"kind": "heading", "level": 2},
            ),
            FormatOperation(
                kind="apply_paragraph",
                target=2,
                parameters={"kind": "body", "level": None},
            ),
            FormatOperation(kind="apply_page", target="all"),
            FormatOperation(kind="format_tables", target="all"),
            FormatOperation(kind="format_images", target="all"),
        ],
    )


def test_formatter_changes_copy_and_preserves_source(tmp_path, profile):
    plan = make_plan(tmp_path)
    before = plan.source.read_bytes()

    output = execute_docx_plan(plan, profile)

    assert plan.source.read_bytes() == before
    assert output != plan.source
    document = Document(output)
    assert document.paragraphs[0].style.name == "Heading 1"
    assert document.paragraphs[1].style.name == "Heading 2"
    assert document.paragraphs[2].style.name == "Normal"
    assert round(document.sections[0].left_margin.cm, 1) == profile.page.left_cm


def test_formatter_repeats_table_header_and_scales_wide_image(tmp_path, profile):
    plan = make_plan(tmp_path)

    output = execute_docx_plan(plan, profile)

    document = Document(output)
    header = document.tables[0].rows[0]._tr.get_or_add_trPr().find(qn("w:tblHeader"))
    assert header is not None
    assert max(shape.width.cm for shape in document.inline_shapes) <= (
        profile.image.max_width_cm + 0.01
    )
