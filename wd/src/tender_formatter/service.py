from pathlib import Path
import hashlib
import tempfile

from docx import Document

from tender_formatter.analyzer import analyze_docx
from tender_formatter.classifier import classify_paragraphs
from tender_formatter.domain import (
    BlockDecision,
    BlockKind,
    DocumentAnalysis,
    FormatProfile,
    ProcessingResult,
)
from tender_formatter.formatter import execute_docx_plan
from tender_formatter.planner import build_plan
from tender_formatter.report import write_report
from tender_formatter.word_adapter import WordAdapter, WordSettings


class FormatterService:
    def __init__(self, word: WordAdapter | None = None):
        self._word = word or WordAdapter()

    def analyze(
        self, source: Path, profile: FormatProfile
    ) -> DocumentAnalysis:
        analysis = analyze_docx(source)
        decisions = classify_paragraphs(analysis.paragraphs, profile)
        return analysis.model_copy(update={"decisions": decisions})

    def format(
        self,
        analysis: DocumentAnalysis,
        profile: FormatProfile,
        overrides: dict[int, BlockDecision],
        output: Path,
        cover_values: dict[str, str],
    ) -> ProcessingResult:
        if output.exists():
            raise ValueError(f"输出文件已存在：{output}")
        if profile.template_path is None or not profile.template_path.is_file():
            raise ValueError("必须选择有效的企业 Word 样板")
        output.parent.mkdir(parents=True, exist_ok=True)
        source_hash = hashlib.sha256(analysis.source.read_bytes()).digest()
        settings = WordSettings(body_page_start=profile.page.body_page_start)
        with tempfile.TemporaryDirectory(
            prefix="tender_formatter_", dir=output.parent
        ) as temporary_directory:
            temporary = Path(temporary_directory)
            content_path = temporary / "formatted_content.docx"
            assembled_path = temporary / "assembled.docx"
            plan = build_plan(analysis, profile, overrides, content_path)
            execute_docx_plan(plan, profile)
            self._word.assemble(
                content_path,
                assembled_path,
                profile.template_path,
                settings,
                cover_values,
            )
            Document(assembled_path)
            if hashlib.sha256(analysis.source.read_bytes()).digest() != source_hash:
                raise RuntimeError("源文件在处理期间发生变化，已停止发布")
            assembled_path.replace(output)

        report_path = output.with_suffix(".report.json")
        warnings = list(plan.warnings) + list(analysis.structure_warnings)
        write_report(
            report_path,
            operation_count=len(plan.operations),
            warnings=warnings,
            paragraph_count=len(analysis.paragraphs),
            heading_count=sum(
                decision.kind == BlockKind.HEADING
                for decision in analysis.decisions
            ),
            confirmed_count=len(overrides),
            template_name=profile.template_path.name,
        )
        return ProcessingResult(
            output=output,
            report=report_path,
            operation_count=len(plan.operations),
            warnings=warnings,
        )
