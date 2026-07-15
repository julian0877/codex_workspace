from pathlib import Path
from zipfile import BadZipFile
import re

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from tender_formatter.domain import DocumentAnalysis, ParagraphEvidence


def _integer_property(element, child_name: str) -> int | None:
    properties = element.pPr
    if properties is None:
        return None
    child = getattr(properties, child_name, None)
    if child is None or child.val is None:
        return None
    return int(child.val)


def _numbering_level(paragraph) -> int | None:
    properties = paragraph._p.pPr
    if properties is None or properties.numPr is None:
        return None
    level = properties.numPr.ilvl
    if level is None or level.val is None:
        return None
    return int(level.val)


def _bold_ratio(paragraph) -> float:
    visible_runs = [(run, len(run.text)) for run in paragraph.runs if run.text]
    total = sum(length for _, length in visible_runs)
    if total == 0:
        return 0.0
    bold = sum(length for run, length in visible_runs if run.bold is True)
    return bold / total


def _extract_paragraph(paragraph) -> ParagraphEvidence:
    alignment = paragraph.alignment
    return ParagraphEvidence(
        text=paragraph.text,
        style_name=paragraph.style.name if paragraph.style is not None else "",
        outline_level=_integer_property(paragraph._p, "outlineLvl"),
        numbering_level=_numbering_level(paragraph),
        bold_ratio=_bold_ratio(paragraph),
        alignment=str(alignment) if alignment is not None else None,
    )


def analyze_docx(path: Path) -> DocumentAnalysis:
    if path.suffix.lower() != ".docx":
        raise ValueError("第一版仅支持 .docx 文件")
    try:
        document = Document(path)
    except (PackageNotFoundError, BadZipFile, KeyError) as exc:
        raise ValueError(f"无法打开 Word 文件：{path.name}") from exc

    paragraphs = [_extract_paragraph(paragraph) for paragraph in document.paragraphs]
    image_count = len(document.part.element.xpath(".//a:blip"))
    warnings: list[str] = []
    caption_count = sum(
        bool(re.match(r"^图\s*\d+(?:[.-]\d+)*\s+", item.text.strip()))
        for item in paragraphs
    )
    if image_count > caption_count:
        warnings.append("存在未配置题注的图片")
    if document.part.element.xpath(".//wp:anchor"):
        warnings.append("存在浮动图片，需要人工复核")
    if any(section.page_width > section.page_height for section in document.sections):
        warnings.append("存在横向页面，分节格式需要人工复核")
    risky_table_indexes: list[int] = []
    for table_index, table in enumerate(document.tables):
        nested = any(cell.tables for row in table.rows for cell in row.cells)
        merged = bool(table._tbl.xpath(".//w:gridSpan | .//w:vMerge"))
        if nested or merged:
            risky_table_indexes.append(table_index)
    if any(
        cell.tables
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    ):
        warnings.append("存在嵌套表格，需要人工复核")
    if document.part.element.xpath(".//w:gridSpan | .//w:vMerge"):
        warnings.append("存在合并单元格，需要人工复核")
    return DocumentAnalysis(
        source=path,
        paragraphs=paragraphs,
        decisions=[],
        table_count=len(document.tables),
        image_count=image_count,
        section_count=len(document.sections),
        structure_warnings=warnings,
        risky_table_indexes=risky_table_indexes,
    )
