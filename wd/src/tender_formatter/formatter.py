import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from tender_formatter.domain import FormatPlan, FormatProfile, TextStyleRules


_PARAGRAPH_ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}
_TABLE_ALIGNMENTS = {
    "left": WD_TABLE_ALIGNMENT.LEFT,
    "center": WD_TABLE_ALIGNMENT.CENTER,
    "right": WD_TABLE_ALIGNMENT.RIGHT,
}


def _apply_style_rules(style, rules: TextStyleRules) -> None:
    style.font.name = rules.latin_font
    style.font.size = Pt(rules.size_pt)
    style.font.bold = rules.bold
    style.element.get_or_add_rPr().get_or_add_rFonts().set(
        qn("w:eastAsia"), rules.east_asia_font
    )
    paragraph_format = style.paragraph_format
    paragraph_format.alignment = _PARAGRAPH_ALIGNMENTS[rules.alignment]
    paragraph_format.line_spacing = rules.line_spacing
    paragraph_format.space_before = Pt(rules.space_before_pt)
    paragraph_format.space_after = Pt(rules.space_after_pt)


def _configure_styles(document, profile: FormatProfile) -> None:
    _apply_style_rules(document.styles["Normal"], profile.body)
    for level in (1, 2, 3):
        rules = profile.headings.get(level, profile.body)
        _apply_style_rules(document.styles[f"Heading {level}"], rules)


def _apply_page_rules(document, profile: FormatProfile) -> None:
    for section in document.sections:
        section.top_margin = Cm(profile.page.top_cm)
        section.bottom_margin = Cm(profile.page.bottom_cm)
        section.left_margin = Cm(profile.page.left_cm)
        section.right_margin = Cm(profile.page.right_cm)
        section.gutter = Cm(profile.page.gutter_cm)


def _repeat_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    header = properties.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        properties.append(header)
    header.set(qn("w:val"), "true")


def _format_tables(
    document, profile: FormatProfile, risky_indexes: set[int]
) -> None:
    for table_index, table in enumerate(document.tables):
        if table_index in risky_indexes:
            continue
        table.alignment = _TABLE_ALIGNMENTS[profile.table.alignment]
        table.autofit = True
        table_width = table._tbl.tblPr.find(qn("w:tblW"))
        if table_width is None:
            table_width = OxmlElement("w:tblW")
            table._tbl.tblPr.insert(0, table_width)
        table_width.set(qn("w:type"), "pct")
        table_width.set(qn("w:w"), str(profile.table.width_percent * 50))
        if profile.table.repeat_header and table.rows:
            _repeat_header(table.rows[0])


def _format_images(document, profile: FormatProfile) -> None:
    maximum_width = Cm(profile.image.max_width_cm)
    for shape in document.inline_shapes:
        if shape.width > maximum_width:
            ratio = maximum_width / shape.width
            shape.width = maximum_width
            shape.height = int(shape.height * ratio)
    for paragraph in document.paragraphs:
        if paragraph._p.xpath(".//a:blip"):
            paragraph.alignment = _PARAGRAPH_ALIGNMENTS[profile.image.alignment]


def execute_docx_plan(plan: FormatPlan, profile: FormatProfile) -> Path:
    if plan.source.resolve() == plan.output.resolve():
        raise ValueError("输出文件不能覆盖源文件")
    plan.output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(plan.source, plan.output)
    document = Document(plan.output)
    _configure_styles(document, profile)

    for operation in plan.operations:
        if operation.kind == "apply_paragraph":
            paragraph = document.paragraphs[int(operation.target)]
            kind = operation.parameters["kind"]
            level = operation.parameters.get("level")
            if kind == "heading":
                paragraph.style = document.styles[f"Heading {level}"]
            elif kind == "caption":
                paragraph.style = document.styles["Caption"]
            else:
                paragraph.style = document.styles["Normal"]
        elif operation.kind == "apply_page":
            _apply_page_rules(document, profile)
        elif operation.kind == "format_tables":
            _format_tables(
                document,
                profile,
                set(operation.parameters.get("risky_indexes", [])),
            )
        elif operation.kind == "format_images":
            _format_images(document, profile)

    document.save(plan.output)
    return plan.output
