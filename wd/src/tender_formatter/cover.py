import re
from collections.abc import Iterable
from pathlib import Path

from docx import Document


_FIELD = re.compile(r"\{\{([^{}]+)\}\}")


def _paragraphs(document) -> Iterable:
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def find_cover_fields(template: Path) -> set[str]:
    document = Document(template)
    return {
        match.group(1).strip()
        for paragraph in _paragraphs(document)
        for match in _FIELD.finditer(paragraph.text)
    }


def _replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def replace_cover_fields(document, values: dict[str, str]) -> list[str]:
    missing: set[str] = set()
    for paragraph in _paragraphs(document):
        original = paragraph.text

        def replace(match: re.Match) -> str:
            field = match.group(1).strip()
            if field not in values or not values[field]:
                missing.add(field)
                return match.group(0)
            return values[field]

        replaced = _FIELD.sub(replace, original)
        if replaced != original:
            _replace_paragraph_text(paragraph, replaced)
    return sorted(missing)
