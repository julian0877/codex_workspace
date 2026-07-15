import math
import tkinter as tk
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from tkinter import filedialog, messagebox, ttk

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


@dataclass
class CheckResult:
    name: str
    actual: float
    allowable: float
    unit: str
    expression: str
    substitution: str = ""
    word_formula: str = ""

    @property
    def ratio(self) -> float:
        if self.allowable == 0:
            return float("inf")
        return self.actual / self.allowable

    @property
    def passed(self) -> bool:
        return self.actual <= self.allowable


MATERIAL_STRENGTHS = {
    "Q235": [
        (16, 215, 125, 320),
        (40, 205, 120, 320),
        (100, 200, 115, 320),
    ],
    "Q345": [
        (16, 305, 175, 400),
        (40, 295, 170, 400),
        (63, 290, 165, 400),
        (80, 280, 160, 400),
        (100, 270, 155, 400),
    ],
    "Q390": [
        (16, 345, 200, 415),
        (40, 330, 190, 415),
        (63, 310, 180, 415),
        (100, 295, 170, 415),
    ],
    "Q420": [
        (16, 375, 215, 440),
        (40, 355, 205, 440),
        (63, 320, 185, 440),
        (100, 305, 175, 440),
    ],
    "40Cr": [
        (100, 365, 210, None),
        (9999, 340, 195, None),
    ],
    "35CrMo": [
        (100, 365, 210, None),
        (9999, 340, 195, None),
    ],
}


TEMPLATES = {
    "5t": {
        "weight": 5,
        "safety_factor": 1.32,
        "lift_points": 2,
        "lug_plates_per_point": 2,
        "angle_deg": 45,
        "pull_force": 44,
        "material": "Q345",
        "plate_thickness": 10,
        "cover_plate_thickness": 0,
        "plate_width": 96,
        "root_gap": 5,
        "plate_height": 88,
        "hole_diameter": 30,
        "weld_beta": 1,
    },
    "10t": {
        "weight": 10,
        "safety_factor": 1.32,
        "lift_points": 2,
        "lug_plates_per_point": 2,
        "angle_deg": 45,
        "pull_force": 0,
        "material": "Q345",
        "plate_thickness": 14,
        "cover_plate_thickness": 0,
        "plate_width": 130,
        "root_gap": 5,
        "plate_height": 120,
        "hole_diameter": 40,
        "weld_beta": 1,
    },
    "15t": {
        "weight": 15,
        "safety_factor": 1.32,
        "lift_points": 2,
        "lug_plates_per_point": 2,
        "angle_deg": 45,
        "pull_force": 0,
        "material": "Q345",
        "plate_thickness": 18,
        "cover_plate_thickness": 0,
        "plate_width": 160,
        "root_gap": 5,
        "plate_height": 150,
        "hole_diameter": 50,
        "weld_beta": 1,
    },
    "30t": {
        "weight": 30,
        "safety_factor": 1.32,
        "lift_points": 2,
        "lug_plates_per_point": 2,
        "angle_deg": 45,
        "pull_force": 0,
        "material": "Q345",
        "plate_thickness": 28,
        "cover_plate_thickness": 0,
        "plate_width": 230,
        "root_gap": 0,
        "plate_height": 210,
        "hole_diameter": 75,
        "weld_beta": 1,
    },
    "40t": {
        "weight": 40,
        "safety_factor": 1.32,
        "lift_points": 2,
        "lug_plates_per_point": 2,
        "angle_deg": 45,
        "pull_force": 0,
        "material": "Q345",
        "plate_thickness": 34,
        "cover_plate_thickness": 0,
        "plate_width": 270,
        "root_gap": 0,
        "plate_height": 250,
        "hole_diameter": 90,
        "weld_beta": 1,
    },
    "50t": {
        "weight": 50,
        "safety_factor": 1.32,
        "lift_points": 2,
        "lug_plates_per_point": 2,
        "angle_deg": 45,
        "pull_force": 0,
        "material": "Q345",
        "plate_thickness": 40,
        "cover_plate_thickness": 0,
        "plate_width": 310,
        "root_gap": 0,
        "plate_height": 290,
        "hole_diameter": 100,
        "weld_beta": 1,
    },
}


def material_strength(material, thickness):
    if material not in MATERIAL_STRENGTHS:
        raise ValueError(f"暂不支持材料：{material}")
    for limit, tensile, shear, bearing in MATERIAL_STRENGTHS[material]:
        if thickness <= limit:
            return {
                "tensile": tensile,
                "shear": shear,
                "bearing": bearing if bearing else 385,
                "weld": 160 if material == "Q235" else 200,
            }
    raise ValueError(f"{material} 厚度超过材料强度表范围，请核对。")


def num(data, key):
    return float(data[key])


def calculate(data):
    weight = num(data, "weight")
    safety_factor = num(data, "safety_factor")
    lift_points = num(data, "lift_points")
    lug_plates_per_point = num(data, "lug_plates_per_point")
    angle_deg = num(data, "angle_deg")
    angle = angle_deg / 180 * math.pi
    t = num(data, "plate_thickness")
    t1 = num(data, "cover_plate_thickness")
    b_width = num(data, "plate_width")
    gap = num(data, "root_gap")
    height = num(data, "plate_height")
    hole_diameter = num(data, "hole_diameter")
    beta = num(data, "weld_beta")
    material = data["material"]

    if lift_points <= 0 or lug_plates_per_point <= 0:
        raise ValueError("吊点数量和每吊点耳板片数必须大于 0。")

    if angle <= 0 or math.sin(angle) <= 0:
        raise ValueError("吊装角度必须大于 0 度。")

    pull_force = num(data, "pull_force")
    if pull_force <= 0:
        pull_force = weight * 9.8 * safety_factor / lift_points / math.sin(angle)

    vertical_force = pull_force * math.sin(angle)
    horizontal_force = pull_force * math.cos(angle)
    plate_force = pull_force / lug_plates_per_point
    plate_vertical_force = vertical_force / lug_plates_per_point
    plate_horizontal_force = horizontal_force / lug_plates_per_point
    weld_size = 0.7 * t
    pin_diameter = hole_diameter - 4
    e = height - 0.5 * b_width
    a = 0.5 * b_width - 0.5 * hole_diameter
    b = a
    z = math.sqrt((a + 0.5 * hole_diameter) ** 2 - (0.5 * hole_diameter) ** 2)
    beff = 2 * t + 16
    net_width = min(beff, b - hole_diameter / 3)
    connection_width = b_width - gap
    strengths = material_strength(material, t)

    if min(pin_diameter, a, z, net_width, connection_width, weld_size) <= 0:
        raise ValueError("存在非正尺寸，请检查孔径、宽度、高度、间隙和厚度。")

    edge_pass = beff <= b and a >= 4 / 3 * beff
    checks = [
        CheckResult(
            "销轴孔净截面抗拉强度",
            plate_force * 1000 / t / net_width,
            strengths["tensile"],
            "N/mm2",
            "单片耳板力*1000/(t*净宽)",
            f"{plate_force:.3f}*1000/({t:.3f}*{net_width:.3f})",
            "σ = N/(t·b₁) ≤ f",
        ),
        CheckResult(
            "耳板端部抗拉劈开强度",
            2 * plate_force * 1000 / t / a,
            strengths["tensile"],
            "N/mm2",
            "2*单片耳板力*1000/(t*a)",
            f"2*{plate_force:.3f}*1000/({t:.3f}*{a:.3f})",
            "σ = 2N/(t·a) ≤ f",
        ),
        CheckResult(
            "耳板抗剪强度",
            plate_force * 1000 / t / z,
            strengths["shear"],
            "N/mm2",
            "单片耳板力*1000/(t*Z)",
            f"{plate_force:.3f}*1000/({t:.3f}*{z:.3f})",
            "τ = N/(t·Z) ≤ fᵥ",
        ),
        CheckResult(
            "耳板端部承压强度",
            plate_force * 1000 / pin_diameter / (t + 2 * t1),
            strengths["bearing"],
            "N/mm2",
            "单片耳板力*1000/(d*(t+2*t1))",
            f"{plate_force:.3f}*1000/({pin_diameter:.3f}*({t:.3f}+2*{t1:.3f}))",
            "σc = N/[d·(t+2t₁)] ≤ fcb",
        ),
    ]

    section_modulus = (t * b_width * b_width - t * gap * gap) / 6
    normal_stress = plate_vertical_force * 1000 / (t * connection_width)
    moment = plate_horizontal_force * e / 1000
    shear_stress = plate_horizontal_force * 1000 / (t * connection_width)
    bending_stress = moment * 1000000 / section_modulus
    combined_stress = ((normal_stress + bending_stress) ** 2 + 3 * shear_stress**2) ** 0.5
    checks.append(
        CheckResult(
            "耳板与构件连接处截面承载力",
            combined_stress,
            1.1 * strengths["tensile"],
            "N/mm2",
            "sqrt((sigma+sigma')^2+3*tau^2)",
            f"sigma={normal_stress:.3f}, sigma'={bending_stress:.3f}, tau={shear_stress:.3f}; "
            f"sqrt(({normal_stress:.3f}+{bending_stress:.3f})^2+3*{shear_stress:.3f}^2)",
            "σeq = √[(σ+σ′)²+3τ²] ≤ 1.1f",
        )
    )

    weld_throat = 0.7 * weld_size
    weld_length = connection_width - 2 * weld_size
    weld_modulus = 2 * 0.7 * weld_size * ((b_width - 2 * weld_size) ** 2 - gap**2) / 6
    if min(weld_throat, weld_length, weld_modulus) <= 0:
        raise ValueError("焊缝计算长度或截面模量为非正值，请调大宽度或减小焊脚。")

    weld_normal = plate_vertical_force * 1000 / (2 * 0.7 * weld_size * weld_length)
    weld_shear = plate_horizontal_force * 1000 / (2 * 0.7 * weld_size * weld_length)
    weld_bending = moment * 1000000 / weld_modulus
    weld_equivalent = ((weld_normal / beta + weld_bending / beta) ** 2 + weld_shear**2) ** 0.5
    checks.append(
        CheckResult(
            "双面角焊缝承载力",
            weld_equivalent,
            strengths["weld"],
            "N/mm2",
            "sqrt(((sigmaN+sigmaM)/beta)^2+tauV^2)",
            f"sigmaN={weld_normal:.3f}, sigmaM={weld_bending:.3f}, tauV={weld_shear:.3f}, beta={beta:.3f}; "
            f"sqrt((({weld_normal:.3f}+{weld_bending:.3f})/{beta:.3f})^2+{weld_shear:.3f}^2)",
            "σf = √[((σN+σM)/βf)²+τV²] ≤ ffw",
        )
    )

    derived = {
        "pull_force": pull_force,
        "vertical_force": vertical_force,
        "horizontal_force": horizontal_force,
        "plate_force": plate_force,
        "plate_vertical_force": plate_vertical_force,
        "plate_horizontal_force": plate_horizontal_force,
        "lift_points": lift_points,
        "lug_plates_per_point": lug_plates_per_point,
        "weld_size": weld_size,
        "pin_diameter": pin_diameter,
        "edge_a": a,
        "edge_b": b,
        "edge_z": z,
        "beff": beff,
        "net_width": net_width,
        "connection_width": connection_width,
        "section_modulus": section_modulus,
        "moment": moment,
        "normal_stress": normal_stress,
        "shear_stress": shear_stress,
        "bending_stress": bending_stress,
        "weld_length": weld_length,
        "weld_modulus": weld_modulus,
        "weld_throat": weld_throat,
        "weld_normal": weld_normal,
        "weld_shear": weld_shear,
        "weld_bending": weld_bending,
        "strengths": strengths,
        "edge_pass": edge_pass,
    }
    all_passed = edge_pass and all(item.passed for item in checks)
    return derived, checks, all_passed


class LiftingLugApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("钢梁板式吊耳验算程序")
        self.geometry("1180x760")
        self.minsize(1080, 680)
        self.vars = {}
        self.last_result = None
        self._build_style()
        self._build_ui()
        self.apply_template("5t")

    def _build_style(self):
        style = ttk.Style(self)
        style.theme_use("clam")
        style.configure("TFrame", background="#f5f7fb")
        style.configure("Card.TFrame", background="#ffffff", relief="solid", borderwidth=1)
        style.configure("TLabel", background="#f5f7fb", font=("Microsoft YaHei UI", 10))
        style.configure("Card.TLabel", background="#ffffff", font=("Microsoft YaHei UI", 10))
        style.configure("Title.TLabel", background="#f5f7fb", font=("Microsoft YaHei UI", 18, "bold"))
        style.configure("Section.TLabel", background="#ffffff", font=("Microsoft YaHei UI", 12, "bold"))
        style.configure("Pass.TLabel", background="#ffffff", foreground="#0f8a43", font=("Microsoft YaHei UI", 16, "bold"))
        style.configure("Fail.TLabel", background="#ffffff", foreground="#b42318", font=("Microsoft YaHei UI", 16, "bold"))
        style.configure("TButton", font=("Microsoft YaHei UI", 10))

    def _build_ui(self):
        self.columnconfigure(0, weight=0)
        self.columnconfigure(1, weight=1)
        self.rowconfigure(1, weight=1)

        header = ttk.Frame(self, padding=(18, 16, 18, 8))
        header.grid(row=0, column=0, columnspan=2, sticky="ew")
        ttk.Label(header, text="钢梁板式吊耳验算程序", style="Title.TLabel").pack(side="left")

        left = ttk.Frame(self, style="Card.TFrame", padding=14)
        left.grid(row=1, column=0, sticky="nsw", padx=(18, 10), pady=(6, 18))

        right = ttk.Frame(self, style="Card.TFrame", padding=14)
        right.grid(row=1, column=1, sticky="nsew", padx=(0, 18), pady=(6, 18))
        right.columnconfigure(0, weight=1)
        right.rowconfigure(2, weight=1)

        ttk.Label(left, text="参数输入", style="Section.TLabel").grid(row=0, column=0, columnspan=3, sticky="w", pady=(0, 10))
        ttk.Label(left, text="吨位模板", style="Card.TLabel").grid(row=1, column=0, sticky="w", pady=5)
        self.template_var = tk.StringVar(value="5t")
        template = ttk.Combobox(left, textvariable=self.template_var, values=list(TEMPLATES), width=16, state="readonly")
        template.grid(row=1, column=1, sticky="ew", pady=5)
        ttk.Button(left, text="套用", command=lambda: self.apply_template(self.template_var.get())).grid(row=1, column=2, padx=(8, 0))

        fields = [
            ("weight", "钢梁自重 G", "t"),
            ("safety_factor", "安全系数", ""),
            ("lift_points", "吊点数量", "个"),
            ("lug_plates_per_point", "每吊点耳板片数", "片"),
            ("angle_deg", "钢丝绳夹角", "deg"),
            ("pull_force", "吊点拉力 T3", "kN，填 0 自动计算"),
            ("plate_thickness", "耳板厚度 t", "mm"),
            ("cover_plate_thickness", "单侧贴板厚度 t1", "mm"),
            ("plate_width", "耳板宽度 B", "mm"),
            ("root_gap", "根部间隙 A", "mm"),
            ("plate_height", "耳板高度 h", "mm"),
            ("hole_diameter", "耳板孔径 d0", "mm"),
            ("weld_beta", "焊缝增大系数 beta", ""),
        ]
        row = 2
        for key, label, unit in fields:
            ttk.Label(left, text=label, style="Card.TLabel").grid(row=row, column=0, sticky="w", pady=5)
            var = tk.StringVar()
            self.vars[key] = var
            ttk.Entry(left, textvariable=var, width=18).grid(row=row, column=1, sticky="ew", pady=5)
            ttk.Label(left, text=unit, style="Card.TLabel").grid(row=row, column=2, sticky="w", padx=(8, 0), pady=5)
            row += 1

        ttk.Label(left, text="耳板材质", style="Card.TLabel").grid(row=row, column=0, sticky="w", pady=5)
        self.vars["material"] = tk.StringVar()
        ttk.Combobox(left, textvariable=self.vars["material"], values=list(MATERIAL_STRENGTHS), width=16, state="readonly").grid(
            row=row, column=1, sticky="ew", pady=5
        )
        row += 1

        actions = ttk.Frame(left, style="Card.TFrame")
        actions.grid(row=row, column=0, columnspan=3, sticky="ew", pady=(14, 0))
        ttk.Button(actions, text="计算", command=self.recalculate).pack(side="left", fill="x", expand=True)
        ttk.Button(actions, text="导出计算书", command=self.export_report).pack(side="left", fill="x", expand=True, padx=(8, 0))

        self.summary = ttk.Label(right, text="等待计算", style="Pass.TLabel")
        self.summary.grid(row=0, column=0, sticky="w")
        self.derived_text = tk.Text(right, height=8, wrap="word", font=("Consolas", 10), relief="flat", background="#f8fafc")
        self.derived_text.grid(row=1, column=0, sticky="ew", pady=(12, 12))

        columns = ("name", "actual", "allowable", "ratio", "result", "expression")
        self.table = ttk.Treeview(right, columns=columns, show="headings", height=14)
        headings = {
            "name": "验算项目",
            "actual": "计算值",
            "allowable": "允许值",
            "ratio": "利用率",
            "result": "结论",
            "expression": "公式",
        }
        widths = {"name": 210, "actual": 110, "allowable": 110, "ratio": 90, "result": 90, "expression": 260}
        for col in columns:
            self.table.heading(col, text=headings[col])
            self.table.column(col, width=widths[col], anchor="w")
        self.table.grid(row=2, column=0, sticky="nsew")
        scrollbar = ttk.Scrollbar(right, orient="vertical", command=self.table.yview)
        scrollbar.grid(row=2, column=1, sticky="ns")
        self.table.configure(yscrollcommand=scrollbar.set)

        self.note = ttk.Label(
            right,
            text="说明：公式按现有 Excel 验算链整理。正式工程使用前，应结合采用规范版本和企业校审要求复核。",
            style="Card.TLabel",
        )
        self.note.grid(row=3, column=0, sticky="w", pady=(12, 0))

    def apply_template(self, name):
        for key, value in TEMPLATES[name].items():
            self.vars[key].set(str(value))
        self.recalculate()

    def read_input(self):
        return {key: var.get().strip() for key, var in self.vars.items()}

    def recalculate(self):
        try:
            data = self.read_input()
            derived, checks, all_passed = calculate(data)
            self.last_result = (data, derived, checks, all_passed)
            self.render_result(data, derived, checks, all_passed)
        except Exception as exc:
            messagebox.showerror("计算错误", str(exc))

    def render_result(self, data, derived, checks, all_passed):
        self.summary.configure(
            text="总体结论：满足要求" if all_passed else "总体结论：不满足，请调整参数",
            style="Pass.TLabel" if all_passed else "Fail.TLabel",
        )
        self.derived_text.configure(state="normal")
        self.derived_text.delete("1.0", "end")
        lines = [
            f"吊点数量: {derived['lift_points']:.0f} 个    每吊点耳板片数: {derived['lug_plates_per_point']:.0f} 片",
            f"吊点 T3 拉力设计值: {derived['pull_force']:.3f} kN",
            f"吊点 T1 竖向分量: {derived['vertical_force']:.3f} kN    吊点 T2 水平分量: {derived['horizontal_force']:.3f} kN",
            f"单片耳板设计力: {derived['plate_force']:.3f} kN    竖向: {derived['plate_vertical_force']:.3f} kN    水平: {derived['plate_horizontal_force']:.3f} kN",
            f"销轴直径 d: {derived['pin_diameter']:.3f} mm    焊脚高度 hf: {derived['weld_size']:.3f} mm",
            f"a: {derived['edge_a']:.3f} mm    b: {derived['edge_b']:.3f} mm    Z: {derived['edge_z']:.3f} mm",
            f"beff: {derived['beff']:.3f} mm    净宽: {derived['net_width']:.3f} mm",
            f"材料强度 f: {derived['strengths']['tensile']} N/mm2    fv: {derived['strengths']['shear']} N/mm2",
            f"边距要求: {'满足' if derived['edge_pass'] else '不满足'}",
        ]
        self.derived_text.insert("end", "\n".join(lines))
        self.derived_text.configure(state="disabled")

        self.table.delete(*self.table.get_children())
        self.table.insert(
            "",
            "end",
            values=("边距要求", f"beff={derived['beff']:.2f}, a={derived['edge_a']:.2f}", "b与4/3beff", "-", "满足" if derived["edge_pass"] else "不满足", "beff<=b 且 a>=4/3*beff"),
        )
        for item in checks:
            self.table.insert(
                "",
                "end",
                values=(
                    item.name,
                    f"{item.actual:.3f} {item.unit}",
                    f"{item.allowable:.3f} {item.unit}",
                    f"{item.ratio:.3f}",
                    "满足" if item.passed else "不满足",
                    item.expression,
                ),
            )

    def export_report(self):
        if not self.last_result:
            self.recalculate()
        if not self.last_result:
            return
        data, derived, checks, all_passed = self.last_result
        default_name = f"吊耳计算书_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        path = filedialog.asksaveasfilename(
            title="导出计算书",
            defaultextension=".docx",
            initialfile=default_name,
            filetypes=[("Word 文档", "*.docx"), ("所有文件", "*.*")],
        )
        if not path:
            return
        self.build_word_report(path, data, derived, checks, all_passed)
        messagebox.showinfo("导出完成", f"已导出：\n{path}")

    def build_word_report(self, path, data, derived, checks, all_passed):
        document = Document()
        self.configure_word_styles(document)
        labels = self.report_labels()
        weight = float(data["weight"])
        safety_factor = float(data["safety_factor"])
        angle_deg = float(data["angle_deg"])
        plate_thickness = float(data["plate_thickness"])
        plate_width = float(data["plate_width"])
        root_gap = float(data["root_gap"])
        plate_height = float(data["plate_height"])
        hole_diameter = float(data["hole_diameter"])
        cover_plate_thickness = float(data["cover_plate_thickness"])
        weld_beta = float(data["weld_beta"])
        material = data["material"]
        results = {item.name: item for item in checks}

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("钢梁板式吊耳验算计算书")
        title_run.bold = True
        title_run.font.size = Pt(18)

        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        conclusion = document.add_paragraph()
        conclusion.add_run("总体结论：").bold = True
        result_run = conclusion.add_run("满足要求" if all_passed else "不满足，请调整参数")
        result_run.bold = True

        document.add_heading("一、计算说明", level=1)
        self.add_body_text(
            document,
            "本计算书用于钢梁板式吊耳吊装验算。钢梁采用两点吊装方式，每个吊点设置双耳板。"
            "钢丝绳与水平面夹角按输入参数考虑，吊装动力及不均匀影响通过安全系数计入。"
        )
        self.add_body_text(
            document,
            "本计算书主要验算内容包括：吊点钢丝绳拉力及单片耳板分担力、耳板几何边距、"
            "销轴孔净截面抗拉强度、耳板端部抗拉劈开强度、耳板抗剪强度、耳板端部承压强度、"
            "耳板与构件连接处截面承载力以及双面角焊缝承载力。"
        )

        document.add_heading("二、计算简图", level=1)
        self.add_lug_diagram(document)
        for line in [
            "G 为钢梁自重；T3 为单吊点钢丝绳拉力设计值；T1 为吊点竖向分量；T2 为吊点水平分量。",
            "B 为耳板宽度；h 为耳板高度；A 为耳板根部间隙；t 为耳板厚度；d0 为耳板孔径；d 为销轴直径。",
            "a、b、Z 为耳板孔周边计算尺寸；hf 为角焊缝焊脚高度。",
        ]:
            self.add_body_text(document, line)

        document.add_heading("三、已知参数", level=1)
        parameter_rows = [(index, label, data[key]) for index, (key, label) in enumerate(labels.items(), start=1)]
        self.add_parameter_table(document, parameter_rows)

        document.add_heading("四、荷载计算", level=1)
        self.add_formula_paragraph(document, "单吊点钢丝绳拉力设计值：", "T3 = G·g·K/(n·sinθ)")
        self.add_body_text(
            document,
            f"代入数值：T3 = {weight:g}×9.8×{safety_factor:g}/"
            f"({derived['lift_points']:.0f}×sin{angle_deg:g}°) = {derived['pull_force']:.3f} kN。"
        )
        self.add_formula_paragraph(document, "吊点竖向分量：", "T1 = T3·sinθ")
        self.add_body_text(document, f"T1 = {derived['pull_force']:.3f}×sin{angle_deg:g}° = {derived['vertical_force']:.3f} kN。")
        self.add_formula_paragraph(document, "吊点水平分量：", "T2 = T3·cosθ")
        self.add_body_text(document, f"T2 = {derived['pull_force']:.3f}×cos{angle_deg:g}° = {derived['horizontal_force']:.3f} kN。")
        self.add_formula_paragraph(document, "单片耳板分担设计力：", "N = T3/m")
        self.add_body_text(
            document,
            f"N = {derived['pull_force']:.3f}/{derived['lug_plates_per_point']:.0f} = {derived['plate_force']:.3f} kN。"
        )
        self.add_body_text(
            document,
            f"单片耳板竖向分量 N1 = {derived['plate_vertical_force']:.3f} kN；"
            f"单片耳板水平分量 N2 = {derived['plate_horizontal_force']:.3f} kN。"
        )

        document.add_heading("五、材料强度设计值", level=1)
        self.add_body_text(document, f"耳板材质为 {material}，耳板厚度 t = {plate_thickness:g} mm。")
        self.add_body_text(document, f"抗拉、抗压和抗弯强度设计值 f = {derived['strengths']['tensile']} N/mm2。")
        self.add_body_text(document, f"抗剪强度设计值 fv = {derived['strengths']['shear']} N/mm2。")
        self.add_body_text(document, f"端面承压强度设计值 fcb = {derived['strengths']['bearing']} N/mm2。")
        self.add_body_text(document, f"角焊缝强度设计值 ffw = {derived['strengths']['weld']} N/mm2。")

        document.add_heading("六、几何参数计算", level=1)
        self.add_formula_paragraph(document, "销轴直径：", "d = d0 - 4")
        self.add_body_text(document, f"d = {hole_diameter:g} - 4 = {derived['pin_diameter']:.3f} mm。")
        self.add_formula_paragraph(document, "耳板孔至两侧净距：", "a = B/2 - d0/2")
        self.add_body_text(document, f"a = {plate_width:g}/2 - {hole_diameter:g}/2 = {derived['edge_a']:.3f} mm；b = a = {derived['edge_b']:.3f} mm。")
        self.add_formula_paragraph(document, "耳板端部计算宽度：", "Z = √[(a+d0/2)² - (d0/2)²]")
        self.add_body_text(document, f"Z = √[({derived['edge_a']:.3f}+{hole_diameter:g}/2)^2 - ({hole_diameter:g}/2)^2] = {derived['edge_z']:.3f} mm。")
        self.add_formula_paragraph(document, "有效宽度：", "beff = 2t + 16")
        self.add_body_text(document, f"beff = 2×{plate_thickness:g} + 16 = {derived['beff']:.3f} mm。")
        self.add_formula_paragraph(document, "销轴孔净截面计算宽度：", "b1 = min(beff, b-d0/3)")
        self.add_body_text(document, f"b1 = min({derived['beff']:.3f}, {derived['edge_b']:.3f}-{hole_diameter:g}/3) = {derived['net_width']:.3f} mm。")
        self.add_body_text(document, f"耳板与构件连接处计算宽度 l = B - A = {plate_width:g} - {root_gap:g} = {derived['connection_width']:.3f} mm。")
        self.add_body_text(document, f"焊脚高度 hf = 0.7t = 0.7×{plate_thickness:g} = {derived['weld_size']:.3f} mm。")
        self.add_body_text(document, f"焊缝计算厚度 he = 0.7hf = {derived['weld_throat']:.3f} mm。")
        self.add_body_text(document, f"焊缝计算长度 lw = l - 2hf = {derived['connection_width']:.3f} - 2×{derived['weld_size']:.3f} = {derived['weld_length']:.3f} mm。")

        document.add_heading("七、边距要求验算", level=1)
        edge_result = "满足" if derived["edge_pass"] else "不满足"
        self.add_check_text(
            document,
            "边距要求",
            "规范公式：",
            f"代入数值：beff = {derived['beff']:.3f} mm，b = {derived['edge_b']:.3f} mm；"
            f"a = {derived['edge_a']:.3f} mm，4/3*beff = {4 / 3 * derived['beff']:.3f} mm。",
            f"结论：{edge_result}。",
            "beff ≤ b，且 a ≥ 4/3·beff",
        )

        section_titles = [
            "八、销轴孔净截面抗拉强度验算",
            "九、耳板端部抗拉劈开强度验算",
            "十、耳板抗剪强度验算",
            "十一、耳板端部承压强度验算",
            "十二、耳板与构件连接处截面承载力验算",
            "十三、双面角焊缝承载力验算",
        ]
        for section_title, item in zip(section_titles, checks):
            result = "满足" if item.passed else "不满足"
            document.add_heading(section_title, level=1)
            if item.name == "耳板与构件连接处截面承载力":
                self.add_body_text(document, f"耳板与构件连接处截面模量 W = {derived['section_modulus']:.3f} mm3。")
                self.add_body_text(document, f"单片耳板竖向分量作用下正应力 sigma = {derived['normal_stress']:.3f} N/mm2。")
                self.add_body_text(document, f"偏心弯矩 M = {derived['moment']:.3f} kN*m。")
                self.add_body_text(document, f"弯矩作用下正应力 sigma' = {derived['bending_stress']:.3f} N/mm2。")
                self.add_body_text(document, f"水平分量作用下剪应力 tau = {derived['shear_stress']:.3f} N/mm2。")
            elif item.name == "双面角焊缝承载力":
                self.add_body_text(document, f"焊缝截面模量 W1 = {derived['weld_modulus']:.3f} mm3。")
                self.add_body_text(document, f"竖向分量作用下焊缝正应力 sigmaN = {derived['weld_normal']:.3f} N/mm2。")
                self.add_body_text(document, f"水平分量作用下焊缝剪应力 tauV = {derived['weld_shear']:.3f} N/mm2。")
                self.add_body_text(document, f"弯矩作用下焊缝正应力 sigmaM = {derived['weld_bending']:.3f} N/mm2。")
            self.add_check_text(
                document,
                item.name,
                "规范公式：",
                f"代入数值：{item.substitution} = {item.actual:.3f} {item.unit}。",
                f"计算值 = {item.actual:.3f} {item.unit}；允许值 = {item.allowable:.3f} {item.unit}；"
                f"利用率 = {item.ratio:.3f}，结论：{result}。",
                item.word_formula or item.expression,
            )

        document.add_heading("十四、验算结果汇总", level=1)
        summary_lines = [f"1. 边距要求：{edge_result}；"]
        for index, item in enumerate(checks, start=2):
            summary_lines.append(f"{index}. {item.name}：{'满足' if item.passed else '不满足'}；")
        for line in summary_lines:
            self.add_numbered_text(document, line)

        document.add_heading("十五、结论及建议", level=1)
        if all_passed:
            self.add_body_text(document, "本吊耳在当前参数条件下，各项验算均满足要求，可按当前参数采用。")
        else:
            self.add_body_text(
                document,
                "本吊耳在当前参数条件下存在不满足项，不宜直接采用，应调整相关尺寸或参数后重新验算。"
            )
            if not derived["edge_pass"]:
                self.add_body_text(
                    document,
                    f"其中边距要求不满足：a = {derived['edge_a']:.3f} mm，"
                    f"4/3·beff = {4 / 3 * derived['beff']:.3f} mm。"
                )
                for line in [
                    "建议增大耳板宽度 B；",
                    "建议增大孔边距 a、b；",
                    "可结合销轴和吊装要求调整耳板孔径 d0；",
                    "优化耳板外形尺寸后应重新进行耳板本体、承压、连接截面及焊缝承载力验算。",
                ]:
                    self.add_numbered_text(document, line)

        note = document.add_paragraph()
        note.add_run("说明：").bold = True
        note.add_run("本计算书按当前程序内置公式生成，正式工程应用前应结合采用规范版本、设计条件和企业校审要求复核。")

        document.save(path)

    @staticmethod
    def add_lug_diagram(document):
        diagram_path = Path(__file__).resolve().parent / "assets" / "lifting_lug_diagram.png"
        if not diagram_path.exists():
            paragraph = document.add_paragraph()
            paragraph.add_run("计算简图：").bold = True
            paragraph.add_run("未找到内置简图文件 assets/lifting_lug_diagram.png。")
            return

        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(diagram_path), width=Inches(6.4))

        caption = document.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.add_run("图 1  吊耳尺寸、受力及焊缝计算简图")

    @staticmethod
    def add_body_text(document, text):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Pt(21)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(text)

    @staticmethod
    def add_formula_paragraph(document, label, equation_text):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Pt(18)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(label)
        LiftingLugApp.add_equation(paragraph, equation_text)

    @staticmethod
    def add_numbered_text(document, text):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Pt(18)
        paragraph.paragraph_format.first_line_indent = Pt(-18)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(text)

    @staticmethod
    def add_parameter_table(document, rows):
        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.autofit = False
        headers = ["序号", "已知参数", "数值"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        for row in rows:
            cells = table.add_row().cells
            cells[0].text = str(row[0])
            cells[1].text = str(row[1])
            cells[2].text = str(row[2])
        LiftingLugApp.set_fixed_table_widths(table, [900, 3900, 3600])
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Microsoft YaHei"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        document.add_paragraph()

    @staticmethod
    def set_fixed_table_widths(table, widths_dxa):
        table_pr = table._tbl.tblPr
        table_width = table_pr.find(qn("w:tblW"))
        if table_width is None:
            table_width = OxmlElement("w:tblW")
            table_pr.append(table_width)
        table_width.set(qn("w:type"), "dxa")
        table_width.set(qn("w:w"), str(sum(widths_dxa)))

        layout = table_pr.find(qn("w:tblLayout"))
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            table_pr.append(layout)
        layout.set(qn("w:type"), "fixed")

        existing_grid = table._tbl.tblGrid
        if existing_grid is not None:
            table._tbl.remove(existing_grid)
        grid = OxmlElement("w:tblGrid")
        for width in widths_dxa:
            grid_col = OxmlElement("w:gridCol")
            grid_col.set(qn("w:w"), str(width))
            grid.append(grid_col)
        table._tbl.insert(1, grid)

        for row in table.rows:
            for index, cell in enumerate(row.cells):
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_width = tc_pr.find(qn("w:tcW"))
                if tc_width is None:
                    tc_width = OxmlElement("w:tcW")
                    tc_pr.append(tc_width)
                tc_width.set(qn("w:type"), "dxa")
                tc_width.set(qn("w:w"), str(widths_dxa[index]))

    @staticmethod
    def add_check_text(document, title, formula_text, substitution_text, result_text, equation_text=None):
        title_paragraph = document.add_paragraph()
        title_paragraph.paragraph_format.space_before = Pt(6)
        title_paragraph.paragraph_format.space_after = Pt(2)
        title_run = title_paragraph.add_run(title)
        title_run.bold = True

        formula_paragraph = document.add_paragraph()
        formula_paragraph.paragraph_format.left_indent = Pt(18)
        formula_paragraph.paragraph_format.space_after = Pt(2)
        formula_paragraph.add_run(formula_text)
        if equation_text:
            LiftingLugApp.add_equation(formula_paragraph, equation_text)

        substitution_paragraph = document.add_paragraph()
        substitution_paragraph.paragraph_format.left_indent = Pt(18)
        substitution_paragraph.paragraph_format.space_after = Pt(2)
        substitution_paragraph.add_run(substitution_text)

        result_paragraph = document.add_paragraph()
        result_paragraph.paragraph_format.left_indent = Pt(18)
        result_paragraph.paragraph_format.space_after = Pt(4)
        result_paragraph.add_run(result_text)

    @staticmethod
    def add_equation(paragraph, equation_text):
        math = OxmlElement("m:oMath")
        math_run = OxmlElement("m:r")
        math_text = OxmlElement("m:t")
        math_text.text = equation_text
        math_run.append(math_text)
        math.append(math_run)
        paragraph._p.append(math)

    @staticmethod
    def configure_word_styles(document):
        for style_name in ["Normal", "Heading 1", "Heading 2"]:
            style = document.styles[style_name]
            style.font.name = "Microsoft YaHei"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        document.styles["Normal"].font.size = Pt(10.5)

    @staticmethod
    def report_labels():
        return {
            "weight": "钢梁自重 G(t)",
            "safety_factor": "安全系数",
            "lift_points": "吊点数量",
            "lug_plates_per_point": "每吊点耳板片数",
            "angle_deg": "钢丝绳夹角(deg)",
            "pull_force": "吊点拉力 T3(kN，0 表示自动)",
            "material": "耳板材质",
            "plate_thickness": "耳板厚度 t(mm)",
            "cover_plate_thickness": "单侧贴板厚度 t1(mm)",
            "plate_width": "耳板宽度 B(mm)",
            "root_gap": "根部间隙 A(mm)",
            "plate_height": "耳板高度 h(mm)",
            "hole_diameter": "耳板孔径 d0(mm)",
            "weld_beta": "焊缝增大系数 beta",
        }

    def build_report(self, data, derived, checks, all_passed):
        lines = [
            "钢梁板式吊耳验算计算书",
            f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            "",
            "一、输入参数",
        ]
        labels = {
            "weight": "钢梁自重 G(t)",
            "safety_factor": "安全系数",
            "lift_points": "吊点数量",
            "lug_plates_per_point": "每吊点耳板片数",
            "angle_deg": "钢丝绳夹角(deg)",
            "pull_force": "吊点拉力 T3(kN，0 表示自动)",
            "material": "耳板材质",
            "plate_thickness": "耳板厚度 t(mm)",
            "cover_plate_thickness": "单侧贴板厚度 t1(mm)",
            "plate_width": "耳板宽度 B(mm)",
            "root_gap": "根部间隙 A(mm)",
            "plate_height": "耳板高度 h(mm)",
            "hole_diameter": "耳板孔径 d0(mm)",
            "weld_beta": "焊缝增大系数 beta",
        }
        for key, label in labels.items():
            lines.append(f"{label}：{data[key]}")
        lines.extend(
            [
                "",
                "二、主要中间量",
                f"吊点数量={derived['lift_points']:.0f}，每吊点耳板片数={derived['lug_plates_per_point']:.0f}",
                f"吊点 T3={derived['pull_force']:.3f} kN，吊点 T1={derived['vertical_force']:.3f} kN，吊点 T2={derived['horizontal_force']:.3f} kN",
                f"单片耳板设计力={derived['plate_force']:.3f} kN，竖向={derived['plate_vertical_force']:.3f} kN，水平={derived['plate_horizontal_force']:.3f} kN",
                f"d={derived['pin_diameter']:.3f} mm，hf={derived['weld_size']:.3f} mm，a={derived['edge_a']:.3f} mm，b={derived['edge_b']:.3f} mm，Z={derived['edge_z']:.3f} mm",
                f"beff={derived['beff']:.3f} mm，净宽={derived['net_width']:.3f} mm，焊缝计算长度={derived['weld_length']:.3f} mm",
                "",
                "三、验算结果",
                f"边距要求：{'满足' if derived['edge_pass'] else '不满足'}",
            ]
        )
        for item in checks:
            lines.append(
                f"{item.name}：计算值 {item.actual:.3f} {item.unit}，允许值 {item.allowable:.3f} {item.unit}，利用率 {item.ratio:.3f}，{'满足' if item.passed else '不满足'}"
            )
        lines.extend(["", f"总体结论：{'满足要求' if all_passed else '不满足，请调整参数'}"])
        return "\n".join(lines)


if __name__ == "__main__":
    app = LiftingLugApp()
    app.mainloop()
