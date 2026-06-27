import os
from datetime import datetime
from pathlib import Path
from typing import Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Inches
from openpyxl import load_workbook


def _read_excel_rows(excel_path: str) -> List[Dict[str, str]]:
    wb = load_workbook(excel_path, data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return []

    headers = [str(cell).strip() if cell is not None else '' for cell in rows[0]]
    data_rows = []
    for row in rows[1:]:
        if not any(cell is not None and str(cell).strip() != '' for cell in row):
            continue
        record = {}
        for idx, header in enumerate(headers):
            value = row[idx] if idx < len(row) else ''
            record[header] = str(value).strip() if value is not None else ''
        data_rows.append(record)
    return data_rows


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)


def _add_paragraph(doc: Document, text: str, indent: bool = False) -> None:
    paragraph = doc.add_paragraph()
    if indent:
        paragraph.paragraph_format.first_line_indent = Inches(0.3)
    paragraph.add_run(text)


def _row_text(row):
    return ' '.join(str(cell).strip() for cell in row if cell is not None and str(cell).strip() != '')


def _find_section_indices(rows):
    markers = {}
    for idx, row in enumerate(rows):
        if not row:
            continue
        first = row[0] if len(row) > 0 else None
        second = row[1] if len(row) > 1 else None
        if isinstance(first, str) and first.startswith('一、'):
            markers['一'] = idx
        elif isinstance(first, str) and first.startswith('二、'):
            markers['二'] = idx
        elif isinstance(first, str) and first.startswith('三、'):
            markers['三'] = idx
        elif isinstance(first, str) and first.startswith('四、'):
            markers['四'] = idx
        elif isinstance(first, str) and first.startswith('五、'):
            markers['五'] = idx
        elif isinstance(second, str) and second.startswith('一、'):
            markers['一'] = idx
        elif isinstance(second, str) and second.startswith('二、'):
            markers['二'] = idx
        elif isinstance(second, str) and second.startswith('三、'):
            markers['三'] = idx
        elif isinstance(second, str) and second.startswith('四、'):
            markers['四'] = idx
        elif isinstance(second, str) and second.startswith('五、'):
            markers['五'] = idx
    return markers


def _extract_section_rows(rows, start_idx, end_idx):
    return [row for row in rows[start_idx:end_idx] if any(cell is not None for cell in row)]


def _extract_key_value_pairs(section_rows):
    pairs = []
    for row in section_rows:
        cells = [cell for cell in row if cell is not None and str(cell).strip() != '']
        if not cells:
            continue
        if len(cells) >= 2:
            if isinstance(cells[0], (int, float)):
                key = f'{cells[0]} {cells[1]}'
                value = ' '.join(str(c) for c in cells[2:])
            else:
                key = str(cells[0])
                value = ' '.join(str(c) for c in cells[1:])
        else:
            key = str(cells[0])
            value = ''
        pairs.append((key, value))
    return pairs


def _add_two_column_table(doc: Document, pairs):
    if not pairs:
        return
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.autofit = True
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '项目'
    hdr_cells[1].text = '数值 / 说明'
    for key, value in pairs:
        row_cells = table.add_row().cells
        row_cells[0].text = str(key)
        row_cells[1].text = str(value)


def generate_word_from_excel(excel_path: str, output_path: str) -> bool:
    if not os.path.exists(excel_path):
        raise FileNotFoundError(f'Excel file not found: {excel_path}')

    wb = load_workbook(excel_path, data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        raise ValueError('Excel file contains no data rows')

    doc = Document()
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('吊耳计算书')
    title_run.font.size = Pt(20)
    title_run.font.bold = True

    doc.add_paragraph('生成时间：' + datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    doc.add_paragraph('来源文件：' + os.path.basename(excel_path))

    markers = _find_section_indices(rows)
    one_idx = markers.get('一', 0)
    two_idx = markers.get('二', len(rows))
    three_idx = markers.get('三', len(rows))
    four_idx = markers.get('四', len(rows))
    five_idx = markers.get('五', len(rows))

    # 计算简图
    _add_heading(doc, '一、计算简图', 1)
    summary_rows = _extract_section_rows(rows, one_idx + 1, two_idx)
    if summary_rows:
        _add_paragraph(doc, '本部分根据 Excel 计算表中的参数说明计算简图，主要涉及吊耳受力形式和关键设计参数。')
        _add_two_column_table(doc, _extract_key_value_pairs(summary_rows))
    else:
        _add_paragraph(doc, '本部分根据 Excel 计算表生成。')

    # 已知条件
    _add_heading(doc, '二、已知条件', 1)
    known_rows = _extract_section_rows(rows, two_idx + 1, three_idx)
    if known_rows:
        _add_paragraph(doc, '已知条件包括吊装荷载、材料参数、几何尺寸和设计系数。')
        _add_two_column_table(doc, _extract_key_value_pairs(known_rows))
    else:
        _add_paragraph(doc, '无可用已知条件数据。')

    # 材质、尺寸及力学性能
    _add_heading(doc, '三、材质、尺寸及力学性能', 1)
    material_rows = _extract_section_rows(rows, three_idx + 1, four_idx)
    if material_rows:
        _add_paragraph(doc, '本节列出材料性质、构件尺寸和力学性能设计值。')
        _add_two_column_table(doc, _extract_key_value_pairs(material_rows))
    else:
        _add_paragraph(doc, '无可用材质和尺寸性能数据。')

    # 耳板承载力计算结果及判定
    _add_heading(doc, '四、耳板承载力计算结果及判定', 1)
    ear_result_rows = _extract_section_rows(rows, four_idx + 1, five_idx)
    if ear_result_rows:
        _add_paragraph(doc, '耳板承载力验算结果如下，包含抗拉、抗剪、承压以及连接处截面承载力判断。')
        _add_two_column_table(doc, _extract_key_value_pairs(ear_result_rows))
    else:
        _add_paragraph(doc, '无耳板承载力验算结果。')

    # 双面角焊缝承载力计算结果及判定
    _add_heading(doc, '五、双面角焊缝承载力计算结果及判定', 1)
    weld_rows = _extract_section_rows(rows, five_idx + 1, len(rows))
    if weld_rows:
        _add_paragraph(doc, '双面角焊缝承载力验算结果如下，包含焊缝强度、剪应力和弯矩效应判定。')
        _add_two_column_table(doc, _extract_key_value_pairs(weld_rows))
    else:
        _add_paragraph(doc, '无角焊缝承载力验算结果。')

    output_dir = os.path.dirname(output_path) or '.'
    os.makedirs(output_dir, exist_ok=True)
    doc.save(output_path)
    return True


def main():
    base_dir = Path(__file__).resolve().parent
    excel_path = base_dir.parent / 'cal' / '附件1-钢梁板式吊耳验算-【5t^F50t】-双吊耳-201906-.xlsx'
    output_path = base_dir / 'output' / '吊耳计算书.docx'
    output_path.parent.mkdir(parents=True, exist_ok=True)
    generate_word_from_excel(str(excel_path), str(output_path))
    print(f'生成完成：{output_path}')


if __name__ == '__main__':
    main()
